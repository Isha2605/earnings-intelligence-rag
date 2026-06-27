import os
from docx import Document

# ── paths ──────────────────────────────────────────────────────────────────
RAW_DIR       = "data/raw"
PROCESSED_DIR = "data/processed"

FILES = {
    2022: "msft_10k_2022.docx",
    2023: "msft_10k_2023.docx",
    2024: "msft_10k_2024.docx",
}

# ── main function ───────────────────────────────────────────────────────────
def parse_docx(filepath):
    doc = Document(filepath)
    content = []

    # extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            content.append(text)

    # extract tables — this is where the numbers live
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_text = cell_text.encode('utf-8', errors='ignore').decode('utf-8')
                    row_data.append(cell_text)
            if row_data:
                # join cells with a tab so numbers stay readable
                content.append("\t".join(row_data))

    return "\n".join(content)


def run():
    # create processed/ folder if it doesn't exist yet
    os.makedirs(PROCESSED_DIR, exist_ok=True)

    for year, filename in FILES.items():
        filepath = os.path.join(RAW_DIR, filename)

        print(f"Parsing {filename}...")
        text = parse_docx(filepath)

        # save extracted text as a plain .txt file
        out_path = os.path.join(PROCESSED_DIR, f"msft_10k_{year}.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)

        print(f"  ✓ Saved to {out_path}  ({len(text):,} characters)")


if __name__ == "__main__":
    run()